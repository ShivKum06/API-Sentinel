from docx import Document

path = r"C:\Users\Shiv\Desktop\INIT' 26\Codes - Copy\API_Sentinel_Judge_Presentation.docx"

doc = Document()
doc.add_heading('API Sentinel – Project Explanation for Judge', 0)

content = [
    'This project is called API Sentinel. It is a demo security system built to protect a mock API from common cyber threats like brute force login attempts, API enumeration, rate abuse, and SQL injection-like payloads.',
    'The system sits in front of the API and inspects every incoming request before allowing it through. It checks things like the request path, method, size, query parameters, request body, client IP, failed login attempts, and repeated access patterns.',
    'Once the request is analyzed, the system gives it a threat score and a risk score. Based on that score, it decides whether to allow, monitor, review, or block the request.',
    'Main idea: The project is designed to show how a real API security gateway works in a simplified form. It is not a full production-grade security product, but it demonstrates the core workflow: capture request, detect suspicious behavior, score the risk, make a decision, save the incident, and show it on a live dashboard.',
    'Folder structure and what each part does: sentinel/main.py is the main server file, sentinel/detection.py handles rule-based threat detection, sentinel/risk.py calculates the final risk score, sentinel/models.py defines the data structures, sentinel/ml_anomaly.py adds anomaly logic, sentinel/ml_features.py prepares the feature data, sentinel/ml_train.py trains the model, sentinel/database.py stores all records, and sentinel/response.py maps the risk to allow or block decisions.',
    'The simulator folder contains demo attackers that generate traffic such as normal requests, enumeration attempts, brute force attacks, and injection payloads. These scripts help the project behave like a real demo environment under attack.',
    'The frontend dashboard displays protected API metrics, threat counts, incident lists, and individual incident details. It also supports override actions such as Allow or Block, and the highlights update immediately after the action is taken.',
    'Audit and logging: the project writes a persistent audit log file called sentinel_audit.log to store the timestamp, incident ID, original action, new action, reviewer, and reason. This makes the system traceable and accountable.',
    'Why this project is impressive: this project combines backend security logic, real-time data capture, SQLite persistence, incident monitoring, dashboard UI, ML-inspired anomaly detection, and override-based review workflows.',
    'In simple words: this project acts like a smart guard in front of an API. It watches every request, checks whether it looks suspicious, decides whether to allow or block it, records what happened, and shows everything on a dashboard.',
    'Final judgment summary: this is a strong demo project because it includes real API protection, attack simulation, security scoring, database persistence, admin override features, live dashboard updates, audit logging, and ML anomaly support.'
]

for paragraph in content:
    doc.add_paragraph(paragraph)

# add a small final section
section = doc.add_paragraph()
section.add_run('Conclusion:')
section.runs[0].bold = True

doc.add_paragraph('This project demonstrates a complete API security workflow in an easy-to-understand format and is suitable for a judge or project presentation.')

doc.save(path)
print(f'Generated valid Word document: {path}')
